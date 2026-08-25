"""
Builds Project_Report.docx.

This is a living document: Section content for a finalized version (e.g.
"Version 1") is never edited once that version is frozen (a copy is saved
into Versions/Version_N/ when the user finalizes it). New work in progress
is appended as its own "Version N (In Progress)" section at the end of this
script / the document, listing changes made since the previous version.

When the user says a version is final:
  1. Rename that section's heading from "(In Progress)" to finalized, add a
     row to the Version History table with the finalization date.
  2. Save a copy of the resulting Project_Report.docx into Versions/Version_N/.
  3. Start a new "Version N+1 (In Progress)" section below it in this script
     for whatever comes next.

Run to regenerate:  python generate_report.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x57)
    return h


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def header_row(table, labels, widths_in=None):
    hdr = table.rows[0].cells
    for i, h in enumerate(labels):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr[i], "1F3B57")
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if widths_in:
        for i, w in enumerate(widths_in):
            table.columns[i].width = Inches(w)


# ============================================================================
# Title page
# ============================================================================
title = doc.add_heading("Maharashtra ETR Estimation System", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Project Report")
r.italic = True
r.font.size = Pt(14)

doc.add_paragraph()
meta = doc.add_table(rows=4, cols=2)
meta.style = "Light Grid Accent 1"
meta_data = [
    ("Report status", "Versions 1-2 finalized; Version 3 in progress"),
    ("Guide", "Dr. S. D. Gorantiwar"),
    ("Prepared by", "Vishal Pandey"),
    ("Last updated", "19 August 2026"),
]
for i, (k, v) in enumerate(meta_data):
    meta.cell(i, 0).text = k
    meta.cell(i, 0).paragraphs[0].runs[0].bold = True
    meta.cell(i, 1).text = v
    set_cell_shading(meta.cell(i, 0), "DCE6F1")
meta.columns[0].width = Inches(2.0)
meta.columns[1].width = Inches(4.5)

note = doc.add_paragraph()
note_run = note.add_run(
    "This report is maintained cumulatively as the project develops. Content for a version is "
    "frozen once that version is finalized (a copy is saved into Versions/Version_N/); new work "
    "is appended below as its own \"(In Progress)\" section until it, too, is finalized."
)
note_run.italic = True
note_run.font.size = Pt(9.5)
note_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_page_break()

# ============================================================================
# VERSION 1 CONTENT - finalized 19 Aug 2026. Do not edit once frozen; a copy
# already lives in Versions/Version_1/Project_Report.docx.
# ============================================================================

add_heading("1. Introduction", level=1)
doc.add_paragraph(
    "Evapotranspiration (ET) is a central quantity in irrigation scheduling and water-balance "
    "studies - it represents the combined loss of water from soil and crop surfaces to the "
    "atmosphere, and is the basis for estimating crop water requirement. This project develops a "
    "software system to estimate the reference evapotranspiration (ETR / ET0) across Maharashtra "
    "state, as an input toward a larger irrigation/crop-water-modelling effort."
)
doc.add_paragraph(
    "The work is being carried out under the guidance of Dr. S. D. Gorantiwar, as part of a "
    "staged assignment: begin with the Hargreaves method at a single location, extend it to a "
    "network of locations across the state, and progress toward the fuller FAO-56 "
    "Penman-Monteith method once the base system is in place."
)

add_heading("2. Objective", level=1)
add_bullets([
    "Estimate daily reference evapotranspiration (ETR) for any location in Maharashtra using the "
    "Hargreaves (1985) method.",
    "Extend single-point estimates to a spatial (state-wide) representation using interpolation "
    "between sampled locations.",
    "Provide the result at multiple administrative scales - state, district, and taluka - not "
    "just at arbitrary points.",
    "Support both a forecast/near-term view (using live weather data) and a long-term historical "
    "archive (1951-2025, using India Meteorological Department gridded data).",
    "Lay the groundwork for the next stage of the assignment: the FAO-56 Penman-Monteith method "
    "and, eventually, integration with soil and farm-boundary data for irrigation modelling.",
])

add_heading("3. Methodology", level=1)

add_heading("3.1 Hargreaves Method", level=2)
p = doc.add_paragraph()
p.add_run("ET0 = 0.0023 x (Tmean + 17.8) x sqrt(Tmax - Tmin) x Ra").bold = True
doc.add_paragraph(
    "Tmax and Tmin are the daily maximum and minimum air temperature (deg C). Ra is the "
    "extraterrestrial radiation (expressed in mm/day equivalent), computed purely from a "
    "location's latitude and the day of year using the standard FAO-56 astronomical formulas - "
    "no external radiation data is required for this term. The Hargreaves method was chosen as "
    "the starting point because it needs only temperature data, unlike the fuller Penman-Monteith "
    "method which also requires humidity, wind, and solar radiation records."
)

add_heading("3.2 Validation", level=2)
doc.add_paragraph(
    "The Hargreaves implementation was validated against a confidential reference dataset "
    "(a worked FAO-56 Penman-Monteith / Hargreaves / Modified Penman calculation for Dapoli "
    "station, 1985, provided by the Guide). For 1 January 1985 at Dapoli (latitude 19 deg 24' N):"
)
vtab = doc.add_table(rows=3, cols=3)
vtab.style = "Light Grid Accent 1"
header_row(vtab, ["Quantity", "This system", "Reference"])
vrows = [("Ra (MJ/m2/day)", "26.185", "26.191"), ("ET0 (mm/day)", "4.1236", "4.1245")]
for i, (a, b, c) in enumerate(vrows, start=1):
    vtab.cell(i, 0).text = a
    vtab.cell(i, 1).text = b
    vtab.cell(i, 2).text = c
doc.add_paragraph()
doc.add_paragraph(
    "Agreement is within approximately 0.02%, confirming the formula implementation is correct."
)

add_heading("3.3 Data Sources", level=2)
add_bullets([
    "Live and forecast weather - Open-Meteo API (free, no key required): daily Tmax/Tmin for any "
    "coordinate, historical archive and up to a 16-day forecast.",
    "Historical weather archive (1951-2025) - India Meteorological Department (IMD) gridded daily "
    "temperature dataset, supplied by the user: 39 fixed grid points (1 deg x 1 deg resolution) "
    "covering Maharashtra, provided as NetCDF files.",
    "Administrative boundaries - official state, district (36), and taluka (358) boundary "
    "shapefiles, supplied by the user, reprojected and simplified for use in the application.",
])

add_heading("3.4 Spatial Interpolation (IDW)", level=2)
doc.add_paragraph(
    "To go from a set of sampled locations to a continuous picture of ETR across the state, "
    "Inverse Distance Weighting (IDW) is used: the estimated value at any location is a "
    "distance-weighted average of nearby sampled points, with closer points weighted more "
    "heavily (weight proportional to 1/distance^2). For the live map, a network of points is "
    "sampled across Maharashtra and interpolated onto a fine grid; for the historical archive, "
    "the same principle is applied via a precomputed weight matrix, so that 75 years of daily "
    "data can be aggregated to district and taluka level in a single matrix operation rather "
    "than processed one day at a time."
)

add_heading("4. Application Features", level=1)
doc.add_paragraph("The system is a Streamlit web application with three tabs, described below.")

add_heading("4.1 Point Lookup", level=2)
add_bullets([
    "Click any location inside Maharashtra to get its daily ETR.",
    "Choose either an upcoming forecast period (up to 16 days) or a historical date range.",
    "The location's district and taluka are identified automatically from the boundary data.",
    "Results are shown as a table and chart, with a CSV download option.",
])

add_heading("4.2 Statewide Map", level=2)
add_bullets([
    "An interactive, zoom-adaptive map: zoomed out shows the state-wide average ETR; zooming in "
    "switches to a district-level choropleth, then to taluka-level as you zoom in further.",
    "Clicking any point at any zoom level shows that exact location's ETR, independent of the "
    "district/taluka shading beneath it.",
    "Works for a chosen date - today, an upcoming forecast date, or a past date.",
    "District boundaries are drawn as bold outlines over the taluka layer so it remains clear "
    "where one district ends and the next begins.",
    "CSV downloads are available for the district/taluka averages table, and for a clicked "
    "point's forecast/history series.",
])

add_heading("4.3 Historical Data (1951-2025)", level=2)
add_bullets([
    "Built from the IMD gridded temperature archive described in section 3.3.",
    "Select State, District, or Taluka level, and (for District/Taluka) the specific unit.",
    "The actually-available date range is shown for the exact selection made (computed from the "
    "real data rather than assumed, so it will correctly reflect any gaps a future data update "
    "might introduce).",
    "Download a CSV of the daily ETR series for the selected unit and date range.",
])

doc.add_paragraph()
fig_p = doc.add_paragraph()
fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_p.add_run().add_picture("Versions/Version_1/_report_fig1.png", width=Inches(5.2))
cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_run = cap.add_run(
    "Figure 1: Illustrative taluka-level ETR output (15 January 2024) as it appeared in Version 1, "
    "with district boundaries overlaid in bold. (Version 2 changed the colour scale and boundary "
    "styling further - see the Version 2 section below.)"
)
cap_run.italic = True
cap_run.font.size = Pt(9.5)

add_heading("5. Technical Architecture", level=1)
add_bullets([
    "app.py - the Streamlit application: page layout, the three tabs, and all user interaction.",
    "etr_core.py - shared, framework-independent logic: the Hargreaves formula, boundary loading, "
    "and single-point weather-fetch helpers.",
    "etr_grid.py - the statewide grid: station sampling, batched weather fetching, IDW "
    "interpolation, and map-colour rendering.",
    "etr_historical.py - builds and serves the 1951-2025 historical archive from the IMD NetCDF "
    "data, aggregated to state/district/taluka level.",
])
doc.add_paragraph(
    "Technology stack: Python, Streamlit (web app framework), Folium + streamlit-folium "
    "(interactive maps), GeoPandas/Shapely (geospatial processing), Pandas/NumPy (data handling), "
    "Requests (API calls), Parquet (efficient storage for the historical archive)."
)

add_heading("6. Version History", level=1)
vh = doc.add_table(rows=1, cols=3)
vh.style = "Light Grid Accent 1"
header_row(vh, ["Version", "Date", "Summary"], widths_in=[0.7, 1.1, 4.7])
vh_rows = [
    ("1", "19 Aug 2026",
     "Base system: Hargreaves method, Point Lookup / Statewide Map / Historical Data (1951-2025) "
     "tabs, validated against reference data. Shown to the Guide as base work."),
    ("2", "19 Aug 2026",
     "Statewide map reliability fix; Digital Elevation Model acquired for Maharashtra (prep for "
     "elevation downscaling); boundary visual hierarchy (state/district/taluka by boldness); "
     "green-yellow-red colour scale; taluka name corruption and typo fixes; official district "
     "renames (Ahilyanagar, Chhatrapati Sambhajinagar, Dharashiv); Point Lookup tab redesigned to "
     "match the Statewide Map's transparent boundary styling. See \"Version 2\" section below for "
     "full details."),
]
for v, d, s in vh_rows:
    row = vh.add_row().cells
    row[0].text = v
    row[1].text = d
    row[2].text = s

add_heading("7. Planned Next Steps (as of Version 1)", level=1)
add_bullets([
    "FAO-56 Penman-Monteith method, using additional inputs (humidity, wind speed, solar "
    "radiation) already available from the Open-Meteo API.",
    "Elevation-based downscaling of the interpolated ETR field, using a 1km-resolution Digital "
    "Elevation Model to account for terrain effects such as the Western Ghats, which flat "
    "distance-based interpolation cannot represent.",
    "Eventually, integration with soil-map and farm-boundary data toward the broader irrigation "
    "model, per the Guide's original assignment.",
])

# ============================================================================
# VERSION 2 (IN PROGRESS) - appended below Version 1, never edited into it.
# Keep adding bullets here as changes happen; do not touch the sections above.
# When the user finalizes Version 2: change the heading below to plain
# "Version 2", update the Version History row's date, save a copy into
# Versions/Version_2/Project_Report.docx, and start a new "Version 3
# (In Progress)" section beneath this one for whatever comes next.
# ============================================================================
doc.add_page_break()
add_heading("Version 2 - Changes Since Version 1", level=1)
doc.add_paragraph(
    "Finalized 19 Aug 2026. The following changes were made after Version 1 was finalized and "
    "are adopted in Version 2. Version 1's content above is unchanged and reflects the system "
    "exactly as it stood at that point."
)

add_heading("Changes adopted in Version 2", level=2)
add_bullets([
    "Statewide map reliability: the live weather fetch for the statewide/district/taluka map now "
    "retries with exponential backoff on rate-limit (HTTP 429) responses from the weather API, "
    "and uses larger request batches to reduce the number of API calls needed - fixing a failure "
    "that occurred at the finest station-spacing setting (0.2 deg).",
    "Digital Elevation Model acquired: a whole-India SRTM DEM (supplied by the user) was clipped "
    "to Maharashtra's boundary, reprojected to latitude/longitude, and resampled to approximately "
    "1km resolution. Verified against known terrain (Western Ghats ridge, Konkan coast, Deccan "
    "plateau). This is preparatory work for elevation-based downscaling of the ETR map (see "
    "Version 1's Planned Next Steps) - not yet integrated into the ETR calculation itself.",
    "Boundary visual hierarchy: at every zoom level in the Statewide Map tab, administrative "
    "boundaries are now layered by boldness so it is clear which line is which - the state "
    "outline is boldest, district outlines medium, and taluka outlines thinnest. Previously all "
    "boundary lines looked the same, making it hard to tell where a district or the state itself "
    "ended.",
    "Colour scale changed: the statewide/district/taluka map now uses a green (low ETR) to "
    "yellow (medium) to red (high) colour gradient, replacing the original single-hue blue scale. "
    "This follows the standard convention used for agricultural water-stress/demand maps, at the "
    "user's request.",
    "Taluka name data-quality fix: the source taluka boundary shapefile contained a systematic "
    "character-encoding corruption affecting 21 taluka names - specifically the talukas that "
    "share their district's name (e.g. \"J<lna\" instead of \"Jalna\", \"N<gpur\" instead of "
    "\"Nagpur\"). The corruption followed a consistent, verifiable substitution pattern "
    "(a mis-encoded as '<', u as '#', i as '\\') and was corrected at the source; two further "
    "unrelated spelling typos (\"Kajrat\" -> \"Karjat\", \"Shrimapur\" -> \"Shrirampur\") were also "
    "corrected after cross-checking against the official Ahilyanagar district government website. "
    "This was a source-data fix, not a full audit of all 358 taluka names against authoritative "
    "records, which would be a separate, larger task.",
    "District renaming to current official names: three Maharashtra districts were officially "
    "renamed by the state government (gazette notifications, 2023-2024) - Ahmednagar to "
    "Ahilyanagar, Aurangabad to Chhatrapati Sambhajinagar, and Osmanabad to Dharashiv. Verified "
    "against news coverage of the official government notifications before applying. The two HQ "
    "talukas that were renamed alongside their districts (Aurangabad and Osmanabad talukas) were "
    "updated to match; the Ahmednagar-area taluka is genuinely named \"Nagar\" in the source data "
    "and government records, so it required no change. All administrative boundary data and the "
    "historical archive (1951-2025) were rebuilt with the corrected/renamed values.",
    "Point Lookup tab redesigned to match the Statewide Map's boundary appearance: the same "
    "zoom-adaptive hierarchy (state outline boldest, district medium, taluka thinnest, shown "
    "progressively as the user zooms in) replaces the previous single on/off taluka mesh - but "
    "with no fill colour at any level, so the underlying map (roads, place names) stays fully "
    "visible, making it easier to visually pick a location to click.",
])

doc.add_paragraph()
fig2_p = doc.add_paragraph()
fig2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig2_p.add_run().add_picture("report_assets/_report_fig2_v2.png", width=Inches(5.2))
cap2 = doc.add_paragraph()
cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap2_run = cap2.add_run(
    "Figure 2: The same output as Figure 1, re-rendered with Version 2's green-yellow-red colour "
    "scale and bold state/district boundary overlay (15 Jan 2024)."
)
cap2_run.italic = True
cap2_run.font.size = Pt(9.5)

# ============================================================================
# VERSION 3 (IN PROGRESS) - appended below Version 2, never edited into it.
# Keep adding bullets here as changes happen; do not touch the sections above.
# When the user finalizes Version 3: change the heading below to plain
# "Version 3", update the Version History row's date, save a copy into
# Versions/Version_3/Project_Report.docx, and start a new "Version 4
# (In Progress)" section beneath this one for whatever comes next.
# ============================================================================
doc.add_page_break()
add_heading("Version 3 (In Progress) - Changes Since Version 2", level=1)
doc.add_paragraph(
    "The following changes have been made after Version 2 was finalized (19 Aug 2026) and are "
    "adopted in Version 3, once finalized. Version 1 and Version 2's content above is unchanged "
    "and reflects the system exactly as it stood at each of those points."
)

add_heading("Changes so far", level=2)
add_bullets([
    "Second data source added to the Historical Data tab: Skymet/PoCRA, alongside IMD. PoCRA "
    "(Project on Climate Resilient Agriculture) operates an automatic weather station (AWS) "
    "network across Maharashtra, run by Skymet, far denser than IMD's 39 fixed grid points "
    "(1,400-2,300+ real stations, growing over time). A new module (etr_skymet.py) aggregates "
    "hourly Tmax/Tmin per station to daily, computes Hargreaves ETR per station, then averages "
    "whichever stations actually reported that day into state/district/taluka values - the same "
    "output schema as the IMD archive, so the app's UI and download logic are shared between "
    "both sources with just a data-source selector added.",
    "Unlike IMD's fixed, always-complete network, PoCRA station coverage genuinely varies: "
    "station count changes year to year, there is a real network-wide gap from Nov 2022 to "
    "mid-2023, and per-unit data availability differs (e.g. some districts' stations only came "
    "online partway through 2022 or 2023). Rather than estimating over these gaps, missing "
    "dates are shown as NA in the table/CSV and as a break in the chart line - and Mumbai City "
    "and Mumbai Suburban districts correctly show no data at all, since PoCRA is an "
    "agricultural monitoring network with no stations in those fully urban districts.",
    "Live Skymet option added to Point Lookup: a third source, \"Skymet live (nearest station)\", "
    "alongside the existing Open-Meteo forecast/historical options. Clicking a point finds the "
    "nearest real PoCRA weather station (within 25 km) and fetches its actual observed data via "
    "the official live API (login + station query, confirmed by direct testing to reject future "
    "dates - it serves observed readings only, not a forecast). Scoped to Point Lookup only, not "
    "the Statewide Map: querying all ~2,300 stations for a single day returns roughly 1 million "
    "raw readings, timed at about 7 minutes to fully page through - impractical for an "
    "on-demand map click, whereas a single station's data returns in under a second. API "
    "credentials are kept out of the codebase entirely, via Streamlit's secrets mechanism.",
])

out_path = r"d:\VIP\1.Task\2.Gorantiwar_Sir\Task_1\Project_Report.docx"
doc.save(out_path)
print("Saved:", out_path)
